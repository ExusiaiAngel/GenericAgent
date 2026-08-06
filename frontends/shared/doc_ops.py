"""文档编辑模块 — 为 GenericAgent 提供 docx/xlsx 读写编辑能力。

使用方式（被 ga.py 中的 do_doc_edit 调用）：
```python
from frontends.shared.doc_ops import doc_read, doc_create, doc_edit, xlsx_read, xlsx_create
```
"""

import os, json
from typing import Optional

# ── docx ──────────────────────────────────────────────────────────

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False
    Document = None

# ── openpyxl ──────────────────────────────────────────────────────

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    HAVE_XLSX = True
except ImportError:
    HAVE_XLSX = False
    openpyxl = None


def _docx_avail():
    if not HAVE_DOCX:
        raise RuntimeError("python-docx 未安装，请执行: pip install python-docx")


def doc_read(path: str) -> str:
    """读取 .docx 文件，返回 Markdown 格式文本。"""
    _docx_avail()
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        style = p.style.name.lower() if p.style else ""
        text = p.text.strip()
        if not text:
            continue
        if "heading" in style or "title" in style:
            try:
                level = int(style.replace("heading ", "").replace("headings ", ""))
                lines.append(f"{'#' * level} {text}")
            except (ValueError, TypeError):
                lines.append(f"# {text}")
        else:
            # 检查格式
            prefix = ""
            for run in p.runs:
                if run.bold:
                    prefix = "**"
                if run.italic:
                    prefix = "_"
            lines.append(f"{prefix}{text}")
    lines.append("")
    # 表格
    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- 表格 {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def doc_create(path: str, content: list, title: str = "") -> str:
    """创建 .docx 文件。

    Args:
        path: 输出路径
        content: 内容结构，每项是 dict:
            {"type": "heading", "text": "...", "level": 1}
            {"type": "paragraph", "text": "...", "bold": false, "italic": false}
            {"type": "table", "headers": [...], "rows": [[...], ...]}
            {"type": "page_break"}
        title: 文档标题（可选）

    Returns:
        生成的 .docx 文件路径
    """
    _docx_avail()
    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    for item in content:
        t = item.get("type", "paragraph")
        if t == "heading":
            doc.add_heading(item["text"], level=item.get("level", 1))
        elif t == "paragraph":
            p = doc.add_paragraph()
            run = p.add_run(item.get("text", ""))
            if item.get("bold"):
                run.bold = True
            if item.get("italic"):
                run.italic = True
            if "font_size" in item:
                run.font.size = Pt(item["font_size"])
            align = item.get("align", "")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif t == "table":
            headers = item.get("headers", [])
            rows = item.get("rows", [])
            table = doc.add_table(rows=1 + len(rows), cols=max(len(headers), 1))
            table.style = "Table Grid"
            # 表头
            for ci, h in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            # 数据行
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = str(val)
        elif t == "page_break":
            doc.add_page_break()

    doc.save(path)
    return os.path.abspath(path)


def doc_edit(path: str, operations: list, output_path: Optional[str] = None) -> str:
    """编辑 .docx 文件。

    Args:
        path: 源文件路径
        operations: 操作列表，每项是 dict:
            {"op": "append_paragraph", "text": "...", "bold": false}
            {"op": "append_heading", "text": "...", "level": 2}
            {"op": "replace_text", "old": "...", "new": "..."}
            {"op": "append_table", "headers": [...], "rows": [[...], ...]}
        output_path: 输出路径（默认覆盖原文件）

    Returns:
        输出文件路径
    """
    _docx_avail()
    doc = Document(path)

    for op in operations:
        kind = op.get("op", "")
        if kind == "append_paragraph":
            p = doc.add_paragraph()
            run = p.add_run(op.get("text", ""))
            if op.get("bold"):
                run.bold = True
            if op.get("italic"):
                run.italic = True
        elif kind == "append_heading":
            doc.add_heading(op["text"], level=op.get("level", 1))
        elif kind == "replace_text":
            old, new = op.get("old", ""), op.get("new", "")
            for p in doc.paragraphs:
                if old in p.text:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            break
                    else:
                        # 跨 run 替换：重建
                        full = p.text.replace(old, new)
                        p.clear()
                        p.add_run(full)
        elif kind == "append_table":
            headers = op.get("headers", [])
            rows = op.get("rows", [])
            table = doc.add_table(rows=1 + len(rows), cols=max(len(headers), 1))
            table.style = "Table Grid"
            for ci, h in enumerate(headers):
                table.rows[0].cells[ci].text = h
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = str(val)

    out = output_path or path
    doc.save(out)
    return os.path.abspath(out)


# ── xlsx ──────────────────────────────────────────────────────────

def _xlsx_avail():
    if not HAVE_XLSX:
        raise RuntimeError("openpyxl 未安装，请执行: pip install openpyxl")


def xlsx_read(path: str, sheet: Optional[str] = None) -> str:
    """读取 .xlsx 文件，返回 Markdown 表格文本。"""
    _xlsx_avail()
    wb: openpyxl.Workbook = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
    lines = [f"# 工作表: {ws.title}"]
    for row in ws.iter_rows(values_only=True):
        vals = [str(v) if v is not None else "" for v in row]
        if any(v.strip() for v in vals):
            lines.append(" | ".join(vals))
    wb.close()
    return "\n".join(lines)


def xlsx_create(path: str, sheets: list) -> str:
    """创建 .xlsx 文件。

    Args:
        path: 输出路径
        sheets: 工作表列表，每项是 dict:
            {"name": "Sheet1", "headers": [...], "rows": [[...], ...]}

    Returns:
        文件路径
    """
    _xlsx_avail()
    if not sheets:
        raise ValueError(
            "xlsx 创建需要至少一个工作表：content 应为 [{\"name\": \"Sheet1\", "
            "\"headers\": [...], \"rows\": [[...], ...]}]"
        )
    wb = openpyxl.Workbook()
    # 删除默认 sheet
    wb.remove(wb.active)

    for si, s in enumerate(sheets):
        ws = wb.create_sheet(title=s.get("name", f"Sheet{si+1}"))
        headers = s.get("headers", [])
        rows = s.get("rows", [])
        if headers:
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True)
        for row in rows:
            ws.append(row)

    wb.save(path)
    return os.path.abspath(path)


# ── 统一入口（供 Agent 调用） ────────────────────────────────────

def dispatch(action: str, params: dict) -> dict:
    """统一调度入口，返回 {"ok": bool, "result": str, "path": str}。"""
    try:
        if action == "read":
            path = params["path"]
            ext = os.path.splitext(path)[1].lower()
            if ext == ".docx":
                text = doc_read(path)
            elif ext in (".xlsx", ".xlsm"):
                text = xlsx_read(path, params.get("sheet"))
            else:
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
            return {"ok": True, "result": text, "path": path}

        elif action == "create":
            path = params["path"]
            ext = os.path.splitext(path)[1].lower()
            content = params.get("content", [])
            title = params.get("title", "")
            if ext == ".docx":
                p = doc_create(path, content, title=title)
            elif ext in (".xlsx", ".xlsm"):
                p = xlsx_create(path, content)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(params.get("text", ""))
                p = os.path.abspath(path)
            return {"ok": True, "result": f"文件已创建: {p}", "path": p}

        elif action == "edit":
            path = params["path"]
            ops = params.get("operations", [])
            out = params.get("output_path")
            ext = os.path.splitext(path)[1].lower()
            if ext == ".docx":
                p = doc_edit(path, ops, output_path=out)
            else:
                raise ValueError(f"不支持编辑 {ext} 文件，仅支持 .docx")
            return {"ok": True, "result": f"文件已编辑: {p}", "path": p}

        else:
            return {"ok": False, "result": f"未知操作: {action}"}
    except Exception as e:
        return {"ok": False, "result": f"错误: {e}"}
